"""
Parser for DOCX documents.
"""

import logging
import os
from typing import List
import docx

from src.parsers.base import BaseParser

logger = logging.getLogger(__name__)


class DOCXParser(BaseParser):
    """Parser for DOCX documents."""
    
    @staticmethod
    def supported_mime_types() -> List[str]:
        """
        Get the list of MIME types supported by this parser.
        
        Returns:
            List of supported MIME types
        """
        return [
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'application/msword'
        ]
    
    @classmethod
    def can_parse(cls, mime_type: str) -> bool:
        """
        Check if this parser can handle the given mime type.
        
        Args:
            mime_type: MIME type of the document
            
        Returns:
            True if this parser can handle the mime type, False otherwise
        """
        return mime_type in cls.supported_mime_types() or 'word' in mime_type.lower()
    
    def parse(self, file_path: str) -> str:
        """
        Parse a DOCX file and extract text content.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Extracted text content
        """
        if not os.path.exists(file_path):
            logger.error(f"DOCX file not found: {file_path}")
            return ""
        
        try:
            text_content = []
            
            # Open the document
            doc = docx.Document(file_path)
            
            # Extract document properties if available
            core_properties = doc.core_properties
            if core_properties:
                if core_properties.title:
                    text_content.append(f"Title: {core_properties.title}")
                if core_properties.author:
                    text_content.append(f"Author: {core_properties.author}")
                if core_properties.created:
                    text_content.append(f"Created: {core_properties.created}")
                text_content.append("")  # Empty line
            
            # Extract all paragraphs
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    text_content.append(para.text)
            
            # Extract tables
            for i, table in enumerate(doc.tables):
                text_content.append(f"\n--- Table {i+1} ---")
                
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    
                    if any(row_text):  # Only append non-empty rows
                        text_content.append(" | ".join(row_text))
                
                text_content.append("")  # Empty line
            
            return "\n".join(text_content)
            
        except Exception as e:
            logger.exception(f"Error parsing DOCX file {file_path}: {e}")
            return f"Error parsing DOCX file: {str(e)}"
